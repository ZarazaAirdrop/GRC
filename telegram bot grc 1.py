from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
from datetime import datetime

TOKEN = "7740133597:AAFyxdxfZ-Dkm3jz1RXSS89z9aCEU8dSkTc"

# Variables globales
datos = {}
estado = {}

# Inicio del bot
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "¡Hola! Soy tu Gestor de Riesgo Cripto.\nPor favor, indica si la recompra es para 'long' o 'short'."
    )
    estado[update.effective_user.id] = "tipo_recompra"

# Procesar mensajes paso a paso
async def procesar_datos(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    texto = update.message.text.strip().lower()

    if estado.get(user_id) == "tipo_recompra":
        if texto in ["long", "short"]:
            datos[user_id] = {"tipo_recompra": texto}
            estado[user_id] = "precio_long"
            await update.message.reply_text("Perfecto, ahora dime el precio LONG.")
        else:
            await update.message.reply_text("Por favor, indica 'long' o 'short'.")
    elif estado.get(user_id) == "precio_long":
        datos[user_id]["precio_long"] = float(texto)
        estado[user_id] = "tokens_long"
        await update.message.reply_text("Dime el valor para tokens long.")
    elif estado.get(user_id) == "tokens_long":
        datos[user_id]["tokens_long"] = float(texto)
        estado[user_id] = "precio_short"
        await update.message.reply_text("Dime el valor para precio short.")
    elif estado.get(user_id) == "precio_short":
        datos[user_id]["precio_short"] = float(texto)
        estado[user_id] = "tokens_short"
        await update.message.reply_text("Dime el valor para tokens short.")
    elif estado.get(user_id) == "tokens_short":
        datos[user_id]["tokens_short"] = float(texto)
        estado[user_id] = "precio_actual"
        await update.message.reply_text("Dime el valor para precio actual.")
    elif estado.get(user_id) == "precio_actual":
        datos[user_id]["precio_actual"] = float(texto)
        estado[user_id] = "precio_recompra"
        await update.message.reply_text("Dime el valor para precio recompra.")
    elif estado.get(user_id) == "precio_recompra":
        datos[user_id]["precio_recompra"] = float(texto)
        estado[user_id] = "tokens_recompra"
        await update.message.reply_text("Dime el valor para tokens recompra.")
    elif estado.get(user_id) == "tokens_recompra":
        datos[user_id]["tokens_recompra"] = float(texto)
        estado[user_id] = "capital_total"
        await update.message.reply_text("Dime el valor para capital total.")
    elif estado.get(user_id) == "capital_total":
        datos[user_id]["capital_total"] = float(texto)
        estado[user_id] = "capital_riesgo"
        await update.message.reply_text("Dime el valor para capital riesgo.")
    elif estado.get(user_id) == "capital_riesgo":
        datos[user_id]["capital_riesgo"] = float(texto)
        estado[user_id] = "niveles_stop_loss"
        await update.message.reply_text("Dime el valor para niveles stop loss.")
    elif estado.get(user_id) == "niveles_stop_loss":
        datos[user_id]["niveles_stop_loss"] = int(texto)
        estado[user_id] = "niveles_take_profit"
        await update.message.reply_text("Dime el valor para niveles take profit.")
    elif estado.get(user_id) == "niveles_take_profit":
        datos[user_id]["niveles_take_profit"] = int(texto)
        estado[user_id] = "porcentaje_take_profit"
        await update.message.reply_text("Dime el valor para porcentaje take profit.")
    elif estado.get(user_id) == "porcentaje_take_profit":
        datos[user_id]["porcentaje_take_profit"] = float(texto)
        estado[user_id] = None
        await calcular_resultados(update, datos[user_id])
    else:
        await update.message.reply_text("Algo salió mal. Intenta de nuevo con /start.")

# Calcular resultados y generar archivo Word
async def calcular_resultados(update: Update, datos: dict):
    tipo_recompra = datos["tipo_recompra"]
    nuevo_precio = (
        (datos["tokens_long"] * datos["precio_long"] + datos["tokens_recompra"] * datos["precio_recompra"])
        / (datos["tokens_long"] + datos["tokens_recompra"])
    )

    niveles_stop_loss = calcular_stop_loss(datos, nuevo_precio, tipo_recompra)
    niveles_take_profit = calcular_take_profit(datos, nuevo_precio, tipo_recompra)

    crear_documento(datos, nuevo_precio, niveles_stop_loss, niveles_take_profit)

    resultados = formatear_resultados(datos, nuevo_precio, niveles_stop_loss, niveles_take_profit)
    await update.message.reply_text(resultados, parse_mode="Markdown")
    await update.message.reply_text("¡Cálculos realizados! Se generó un archivo Word con los resultados.")

# Cálculo de Stop Loss y Take Profit
def calcular_stop_loss(datos, nuevo_precio, tipo_recompra):
    niveles = datos["niveles_stop_loss"]
    tokens_recompra = datos["tokens_recompra"]
    factor = 1 if tipo_recompra == "short" else -1
    return [
        {
            "Nivel": i + 1,
            "Precio": nuevo_precio + factor * (0.01 * (i + 1) * nuevo_precio),
            "Tokens": tokens_recompra * ((niveles - i) / niveles)
        }
        for i in range(niveles)
    ]

def calcular_take_profit(datos, nuevo_precio, tipo_recompra):
    niveles = datos["niveles_take_profit"]
    tokens_recompra = datos["tokens_recompra"]
    porcentaje_tp = datos["porcentaje_take_profit"] / 100
    factor = -1 if tipo_recompra == "short" else 1
    return [
        {
            "Nivel": i + 1,
            "Precio": nuevo_precio + factor * (porcentaje_tp * (i + 1) * nuevo_precio),
            "Tokens": tokens_recompra * ((i + 1) / niveles)
        }
        for i in range(niveles)
    ]

def crear_documento(datos, nuevo_precio, niveles_stop_loss, niveles_take_profit):
    doc = Document()
    doc.add_heading("Gestor de Riesgo Cripto (GRC)", level=1)
    doc.add_heading("Datos Ingresados:", level=2)
    for clave, valor in datos.items():
        doc.add_paragraph(f"{clave.capitalize()}: {valor}")
    doc.add_heading("Resultados Calculados:", level=2)
    doc.add_paragraph(f"Nuevo Precio Promedio: {nuevo_precio:.6f}")
    doc.add_heading("Niveles de Stop Loss:", level=2)
    for nivel in niveles_stop_loss:
        doc.add_paragraph(f"Nivel {nivel['Nivel']}: Precio {nivel['Precio']:.6f}, Tokens {nivel['Tokens']:.6f}")
    doc.add_heading("Niveles de Take Profit:", level=2)
    for nivel in niveles_take_profit:
        doc.add_paragraph(f"Nivel {nivel['Nivel']}: Precio {nivel['Precio']:.6f}, Tokens {nivel['Tokens']:.6f}")
    nombre_archivo = f"Resultados/GRC_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nombre_archivo)

def formatear_resultados(datos, nuevo_precio, niveles_stop_loss, niveles_take_profit):
    resultados = "**Datos Ingresados:**\n"
    for clave, valor in datos.items():
        resultados += f"- {clave.capitalize()}: {valor}\n"
    resultados += f"\n**Nuevo Precio Promedio:** {nuevo_precio:.6f}\n\n"
    resultados += "**Niveles de Stop Loss:**\nNivel | Precio    | Tokens\n"
    for nivel in niveles_stop_loss:
        resultados += f"{nivel['Nivel']}     | {nivel['Precio']:.6f} | {nivel['Tokens']:.6f}\n"
    resultados += "\n**Niveles de Take Profit:**\nNivel | Precio    | Tokens\n"
    for nivel in niveles_take_profit:
        resultados += f"{nivel['Nivel']}     | {nivel['Precio']:.6f} | {nivel['Tokens']:.6f}\n"
    return resultados

# Configurar bot
app = ApplicationBuilder().token(TOKEN).build()
app.add_handler(CommandHandler("start", start))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, procesar_datos))

print("Bot en ejecución...")
app.run_polling()

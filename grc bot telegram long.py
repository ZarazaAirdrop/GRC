from telegram import Update, InputFile
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document

# Token del bot de Telegram
BOT_TOKEN = "7740133597:AAFyxdxfZ-Dkm3jz1RXSS89z9aCEU8dSkTc"

# Almacenar datos temporales por usuario
user_data = {}

# Inicia el flujo de datos
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.effective_user.id
    user_data[user_id] = {
        "estado": 0,
        "datos": {}
    }
    await update.message.reply_text("¡Hola! Vamos a configurar tu gestión. Por favor, envíame el precio LONG:")

# Manejar mensajes de datos
async def procesar_datos(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.effective_user.id
    mensaje = update.message.text

    if user_id not in user_data:
        await update.message.reply_text("Por favor, escribe /start para comenzar.")
        return

    estado = user_data[user_id]["estado"]
    datos = user_data[user_id]["datos"]

    try:
        if estado == 0:
            datos["precio_long"] = float(mensaje)
            await update.message.reply_text("Ahora envíame la cantidad de tokens LONG:")
        elif estado == 1:
            datos["tokens_long"] = float(mensaje)
            await update.message.reply_text("Ahora envíame el precio SHORT:")
        elif estado == 2:
            datos["precio_short"] = float(mensaje)
            await update.message.reply_text("Ahora envíame la cantidad de tokens SHORT:")
        elif estado == 3:
            datos["tokens_short"] = float(mensaje)
            await update.message.reply_text("Ahora envíame el precio actual:")
        elif estado == 4:
            datos["precio_actual"] = float(mensaje)
            await update.message.reply_text("Ahora envíame el precio de recompra:")
        elif estado == 5:
            datos["precio_recompra"] = float(mensaje)
            await update.message.reply_text("Ahora envíame la cantidad de tokens para la recompra:")
        elif estado == 6:
            datos["tokens_recompra"] = float(mensaje)
            await update.message.reply_text("Ahora envíame el capital total (USD):")
        elif estado == 7:
            datos["capital_total"] = float(mensaje)
            await update.message.reply_text("Ahora envíame el capital de riesgo (USD):")
        elif estado == 8:
            datos["capital_riesgo"] = float(mensaje)
            await update.message.reply_text("Ahora envíame el número de niveles de Stop Loss:")
        elif estado == 9:
            datos["niveles_stop_loss"] = int(mensaje)
            await update.message.reply_text("Ahora envíame el número de niveles de Take Profit:")
        elif estado == 10:
            datos["niveles_take_profit"] = int(mensaje)
            await update.message.reply_text("Por último, envíame el porcentaje de Take Profit:")
        elif estado == 11:
            datos["porcentaje_take_profit"] = float(mensaje)

            # Calcular resultados
            resultados = calcular_resultados(datos)

            # Generar el documento Word
            archivo_word = generar_documento(datos, resultados)

            # Enviar el documento al usuario
            with open(archivo_word, "rb") as file:
                await update.message.reply_document(document=InputFile(file), filename=archivo_word)

            # Resetear el flujo
            del user_data[user_id]
            return

        # Avanzar al siguiente estado
        user_data[user_id]["estado"] += 1

    except ValueError:
        await update.message.reply_text("Por favor, envía un valor válido.")

# Función para calcular resultados
def calcular_resultados(datos):
    # Calcular el nuevo precio promedio
    nuevo_precio_promedio = (
        (datos["precio_long"] * datos["tokens_long"] +
         datos["precio_recompra"] * datos["tokens_recompra"]) /
        (datos["tokens_long"] + datos["tokens_recompra"])
    )

    # Calcular niveles de Stop Loss
    niveles_stop_loss = []
    for i in range(1, datos["niveles_stop_loss"] + 1):
        precio_stop = nuevo_precio_promedio - i * (datos["capital_riesgo"] / datos["tokens_recompra"])
        resultado_stop = (precio_stop - nuevo_precio_promedio) * datos["tokens_recompra"]
        niveles_stop_loss.append((i, precio_stop, resultado_stop))

    # Calcular niveles de Take Profit
    niveles_take_profit = []
    for i in range(1, datos["niveles_take_profit"] + 1):
        precio_profit = nuevo_precio_promedio + i * (nuevo_precio_promedio * datos["porcentaje_take_profit"] / 100)
        resultado_profit = (precio_profit - nuevo_precio_promedio) * datos["tokens_long"]
        niveles_take_profit.append((i, precio_profit, resultado_profit))

    return {
        "nuevo_precio_promedio": nuevo_precio_promedio,
        "niveles_stop_loss": niveles_stop_loss,
        "niveles_take_profit": niveles_take_profit
    }

# Función para generar el documento Word
def generar_documento(datos, resultados):
    doc = Document()
    doc.add_heading("Gestor de Riesgo Cripto (GRC)", level=1)

    doc.add_heading("Datos Ingresados:", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Dato"
    hdr_cells[1].text = "Valor"

    for clave, valor in datos.items():
        fila = table.add_row().cells
        fila[0].text = clave
        fila[1].text = str(valor)

    doc.add_heading("Resultados:", level=2)
    doc.add_paragraph(f"Nuevo Precio Promedio: {resultados['nuevo_precio_promedio']:.6f}")

    doc.add_heading("Niveles de Stop Loss:", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Nivel"
    hdr_cells[1].text = "Precio"
    hdr_cells[2].text = "Resultado"

    for nivel, precio, resultado in resultados["niveles_stop_loss"]:
        fila = table.add_row().cells
        fila[0].text = str(nivel)
        fila[1].text = f"{precio:.6f}"
        fila[2].text = f"{resultado:.6f}"

    doc.add_heading("Niveles de Take Profit:", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Nivel"
    hdr_cells[1].text = "Precio"
    hdr_cells[2].text = "Resultado"

    for nivel, precio, resultado in resultados["niveles_take_profit"]:
        fila = table.add_row().cells
        fila[0].text = str(nivel)
        fila[1].text = f"{precio:.6f}"
        fila[2].text = f"{resultado:.6f}"

    nombre_archivo = "Gestor_Riesgo_Cripto.docx"
    doc.save(nombre_archivo)
    return nombre_archivo

# Configuración del bot
def main():
    app = Application.builder().token(BOT_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, procesar_datos))

    print("Bot en ejecución...")
    app.run_polling()

if __name__ == "__main__":
    main()
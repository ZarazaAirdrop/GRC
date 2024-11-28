import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def generar_documento():
    # Cargar los datos desde el archivo JSON
    with open("datos_grc.json", "r") as archivo_json:
        datos = json.load(archivo_json)

    # Crear el documento Word
    doc = Document()
    doc.add_heading("Gestor de Riesgo Cripto (GRC)", level=1)

    # Añadir los datos ingresados
    doc.add_heading("Datos Ingresados:", level=2)
    tabla_datos = doc.add_table(rows=1, cols=2)
    tabla_datos.style = "Table Grid"
    tabla_datos.cell(0, 0).text = "Dato"
    tabla_datos.cell(0, 1).text = "Valor"

    for clave, valor in datos["Datos Ingresados"].items():
        fila = tabla_datos.add_row().cells
        fila[0].text = clave
        fila[1].text = str(valor)

    # Añadir el nuevo precio promedio
    doc.add_paragraph(f"Nuevo Precio Promedio: {datos['Nuevo Precio Promedio']:.6f}")

    # Añadir niveles de Stop Loss
    doc.add_heading("Niveles de Stop Loss:", level=2)
    tabla_stop = doc.add_table(rows=1, cols=4)
    tabla_stop.style = "Table Grid"
    encabezados = ["Nivel", "Precio", "Tokens", "Resultado"]
    for i, encabezado in enumerate(encabezados):
        tabla_stop.cell(0, i).text = encabezado

    for nivel in datos["Niveles de Stop Loss"]:
        fila = tabla_stop.add_row().cells
        fila[0].text = str(nivel["Nivel"])
        fila[1].text = f"{nivel['Precio']:.6f}"
        fila[2].text = f"{nivel['Tokens']:.2f}"
        fila[3].text = f"{nivel['Resultado']:.6f}"

    # Añadir niveles de Take Profit
    doc.add_heading("Niveles de Take Profit:", level=2)
    tabla_tp = doc.add_table(rows=1, cols=4)
    tabla_tp.style = "Table Grid"
    for i, encabezado in enumerate(encabezados):
        tabla_tp.cell(0, i).text = encabezado

    for nivel in datos["Niveles de Take Profit"]:
        fila = tabla_tp.add_row().cells
        fila[0].text = str(nivel["Nivel"])
        fila[1].text = f"{nivel['Precio']:.6f}"
        fila[2].text = f"{nivel['Tokens']:.2f}"
        fila[3].text = f"{nivel['Resultado']:.6f}"

    # Guardar el documento
    doc.save("Gestor_Riesgo_Cripto.docx")
    print("\nArchivo 'Gestor_Riesgo_Cripto.docx' generado exitosamente.")

# Ejecutar la generación del documento
generar_documento()

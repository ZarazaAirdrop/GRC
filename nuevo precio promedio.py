from docx import Document
from docx.shared import Pt, RGBColor

def exportar_resultados_word(nombre_archivo, datos, stop_loss, take_profit):
    doc = Document()

    # Título principal
    titulo = doc.add_heading("Gestor de Riesgo Cripto (GRC)", level=1)
    titulo.alignment = 1  # Centrado

    # Datos ingresados
    doc.add_heading("Datos Ingresados:", level=2)
    tabla_datos = doc.add_table(rows=1, cols=2)
    tabla_datos.style = 'Table Grid'
    tabla_datos.autofit = True
    hdr_cells = tabla_datos.rows[0].cells
    hdr_cells[0].text = 'Dato'
    hdr_cells[1].text = 'Valor'
    for key, value in datos.items():
        row_cells = tabla_datos.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

    # Niveles de Stop Loss
    doc.add_heading("Niveles de Stop Loss:", level=2)
    tabla_stop_loss = doc.add_table(rows=1, cols=4)
    tabla_stop_loss.style = 'Table Grid'
    hdr_cells = tabla_stop_loss.rows[0].cells
    hdr_cells[0].text = 'Nivel'
    hdr_cells[1].text = 'Precio'
    hdr_cells[2].text = 'Tokens'
    hdr_cells[3].text = 'Resultado'

    for nivel in stop_loss:
        row_cells = tabla_stop_loss.add_row().cells
        row_cells[0].text = str(nivel['Nivel'])
        row_cells[1].text = f"{nivel['Precio']:.6f}"
        row_cells[2].text = f"{nivel['Tokens']:.1f}"
        row_cells[3].text = f"{nivel['Resultado']:.6f}"

        # Colorear celdas
        for cell in row_cells:
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 0, 0)  # Rojo

    # Niveles de Take Profit
    doc.add_heading("Niveles de Take Profit:", level=2)
    tabla_take_profit = doc.add_table(rows=1, cols=4)
    tabla_take_profit.style = 'Table Grid'
    hdr_cells = tabla_take_profit.rows[0].cells
    hdr_cells[0].text = 'Nivel'
    hdr_cells[1].text = 'Precio'
    hdr_cells[2].text = 'Tokens'
    hdr_cells[3].text = 'Resultado'

    for nivel in take_profit:
        row_cells = tabla_take_profit.add_row().cells
        row_cells[0].text = str(nivel['Nivel'])
        row_cells[1].text = f"{nivel['Precio']:.6f}"
        row_cells[2].text = f"{nivel['Tokens']:.1f}"
        row_cells[3].text = f"{nivel['Resultado']:.6f}"

        # Colorear celdas
        for cell in row_cells:
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(0, 128, 0)  # Verde

    # Guardar archivo
    doc.save(nombre_archivo)
    print(f"Resultados exportados a '{nombre_archivo}' en formato Word.")
